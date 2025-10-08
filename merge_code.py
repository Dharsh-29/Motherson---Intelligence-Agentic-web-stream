from docx import Document
from pathlib import Path
import os
import subprocess

# Create a new Word document
doc = Document()

# Define project root
project_root = Path(".")  # run from project root folder

# Define file extensions to include
file_extensions = [".py", ".txt", ".sh"]

# Loop through files recursively
for file_path in project_root.rglob("*"):
    # Skip .venv, __pycache__, and directories
    if any(part in str(file_path).split(os.sep) for part in [".venv", "__pycache__"]):
        continue

    if file_path.suffix.lower() in file_extensions:
        try:
            # Add file name as heading
            doc.add_heading(str(file_path), level=2)
            # Read file content
            content = file_path.read_text(encoding="utf-8")
            doc.add_paragraph(content)
            # Add a line break
            doc.add_paragraph("\n")
        except Exception as e:
            print(f"Skipping {file_path} due to error: {e}")

# Save Word document
output_file = "All_Code.docx"
doc.save(output_file)
print(f"✅ All code merged into {output_file}")

# Open Word file automatically (Windows)
try:
    subprocess.run(["start", output_file], shell=True)
except Exception as e:
    print(f"Could not open file automatically: {e}")
